"""
Resume Parser Module

Parses resumes from PDF, DOCX, and TXT files to extract structured information
including skills, education, experience, certifications, and contact details.
"""

import re
import os
from typing import Dict, List, Optional, Any
from dataclasses import dataclass, field
from pathlib import Path

# PDF parsing
try:
    import pdfplumber
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    from PyPDF2 import PdfReader
    PYPDF2_AVAILABLE = True
except ImportError:
    PYPDF2_AVAILABLE = False

# DOCX parsing
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


@dataclass
class ParsedResume:
    """Structured representation of a parsed resume."""
    raw_text: str = ""
    name: str = ""
    email: str = ""
    phone: str = ""
    location: str = ""
    linkedin: str = ""
    github: str = ""
    skills: List[str] = field(default_factory=list)
    education: List[Dict[str, str]] = field(default_factory=list)
    experience: List[Dict[str, str]] = field(default_factory=list)
    certifications: List[Dict[str, str]] = field(default_factory=list)
    projects: List[Dict[str, str]] = field(default_factory=list)
    publications: List[Dict[str, str]] = field(default_factory=list)
    awards: List[str] = field(default_factory=list)
    interests: List[str] = field(default_factory=list)
    languages: List[str] = field(default_factory=list)
    summary: str = ""
    file_name: str = ""
    total_years_experience: float = 0.0
    
    def to_text_for_embedding(self) -> str:
        """
        Generate a text representation optimized for embedding generation.
        Includes all sections for comprehensive semantic matching.
        
        Returns:
            Concatenated text of all resume sections
        """
        parts = []
        
        if self.summary:
            parts.append(f"Summary: {self.summary}")
        
        if self.skills:
            parts.append(f"Skills: {', '.join(self.skills)}")
        
        if self.experience:
            exp_texts = []
            for exp in self.experience:
                exp_text = f"{exp.get('title', '')} at {exp.get('company', '')}"
                if exp.get('duration'):
                    exp_text += f" ({exp['duration']})"
                if exp.get('description'):
                    exp_text += f": {exp['description']}"
                if exp.get('achievements'):
                    achievements = exp['achievements']
                    if isinstance(achievements, list):
                        exp_text += f". Key achievements: {'; '.join(achievements[:3])}"
                exp_texts.append(exp_text)
            parts.append(f"Experience: {'; '.join(exp_texts)}")
        
        if self.projects:
            proj_texts = []
            for proj in self.projects:
                proj_text = proj.get('name', '')
                if proj.get('description'):
                    proj_text += f": {proj['description']}"
                if proj.get('technologies'):
                    techs = proj['technologies']
                    if isinstance(techs, list):
                        proj_text += f" (Technologies: {', '.join(techs)})"
                proj_texts.append(proj_text)
            parts.append(f"Projects: {'; '.join(proj_texts)}")
        
        if self.education:
            edu_texts = []
            for edu in self.education:
                edu_text = f"{edu.get('degree', '')} in {edu.get('field', '')} from {edu.get('institution', '')}"
                if edu.get('year'):
                    edu_text += f" ({edu['year']})"
                if edu.get('gpa'):
                    edu_text += f", GPA: {edu['gpa']}"
                edu_texts.append(edu_text)
            parts.append(f"Education: {'; '.join(edu_texts)}")
        
        if self.certifications:
            cert_texts = []
            for cert in self.certifications:
                if isinstance(cert, dict):
                    cert_text = cert.get('name', '')
                    if cert.get('issuer'):
                        cert_text += f" by {cert['issuer']}"
                    if cert.get('year'):
                        cert_text += f" ({cert['year']})"
                else:
                    cert_text = str(cert)
                cert_texts.append(cert_text)
            parts.append(f"Certifications: {', '.join(cert_texts)}")
        
        if self.publications:
            pub_texts = []
            for pub in self.publications:
                if isinstance(pub, dict):
                    pub_text = pub.get('title', '')
                    if pub.get('venue'):
                        pub_text += f" - {pub['venue']}"
                else:
                    pub_text = str(pub)
                pub_texts.append(pub_text)
            parts.append(f"Publications: {'; '.join(pub_texts)}")
        
        if self.awards:
            parts.append(f"Awards: {', '.join(self.awards)}")
        
        if self.interests:
            parts.append(f"Interests: {', '.join(self.interests)}")
        
        if self.languages:
            parts.append(f"Languages: {', '.join(self.languages)}")
        
        # Include raw text if structured extraction is limited
        if len(parts) < 3 and self.raw_text:
            parts.append(self.raw_text[:2000])
        
        return "\n".join(parts)
    
    def to_dict(self) -> Dict[str, Any]:
        """Convert to dictionary representation."""
        return {
            "name": self.name,
            "email": self.email,
            "phone": self.phone,
            "location": self.location,
            "linkedin": self.linkedin,
            "github": self.github,
            "skills": self.skills,
            "education": self.education,
            "experience": self.experience,
            "certifications": self.certifications,
            "projects": self.projects,
            "publications": self.publications,
            "awards": self.awards,
            "interests": self.interests,
            "languages": self.languages,
            "summary": self.summary,
            "file_name": self.file_name,
            "total_years_experience": self.total_years_experience,
            "raw_text": self.raw_text[:500] + "..." if len(self.raw_text) > 500 else self.raw_text
        }


class ResumeParser:
    """
    Parser for extracting structured information from resume files.
    
    Supports PDF, DOCX, and TXT file formats.
    """
    
    # Common skills keywords for extraction
    COMMON_SKILLS = [
        # Programming Languages
        "python", "java", "javascript", "typescript", "c++", "c#", "ruby", "go", "rust",
        "php", "swift", "kotlin", "scala", "r", "matlab", "sql", "html", "css",
        # Frameworks & Libraries
        "react", "angular", "vue", "node.js", "django", "flask", "fastapi", "spring",
        "tensorflow", "pytorch", "keras", "scikit-learn", "pandas", "numpy",
        "langchain", "langgraph", "streamlit",
        # Cloud & DevOps
        "aws", "azure", "gcp", "docker", "kubernetes", "jenkins", "terraform",
        "ci/cd", "git", "github", "gitlab", "linux", "unix",
        # Data & AI
        "machine learning", "deep learning", "nlp", "computer vision", "data science",
        "data analysis", "data engineering", "etl", "big data", "spark", "hadoop",
        "tableau", "power bi", "excel", "statistics",
        # Databases
        "mysql", "postgresql", "mongodb", "redis", "elasticsearch", "oracle",
        "sql server", "dynamodb", "cassandra", "neo4j",
        # Soft Skills
        "leadership", "communication", "teamwork", "problem solving", "agile", "scrum",
        "project management", "analytical", "critical thinking",
    ]
    
    # Section headers for parsing - use \n to anchor to line start
    SECTION_PATTERNS = {
        "experience": r"(?im)^(?:work\s*experience|professional\s*experience|employment\s*history|work\s*history|employment)\s*$",
        "education": r"(?im)^(?:education|academic\s*background|qualifications|degrees?)\s*$",
        "skills": r"(?im)^(?:skills|technical\s*skills|competencies|expertise|technologies|core\s*competencies)\s*$",
        "certifications": r"(?im)^(?:certifications?|certificates?|licenses?\s*(?:&|and)?\s*certifications?|credentials)\s*$",
        "summary": r"(?im)^(?:summary|objective|profile|about\s*me|professional\s*summary)\s*$",
        "projects": r"(?im)^(?:projects?|portfolio|personal\s*projects?|side\s*projects?|publications?\s*(?:&|and)?\s*projects?)\s*$",
        "publications": r"(?im)^(?:publications?|papers?|research)\s*$",
        "awards": r"(?im)^(?:awards?|honors?|achievements?|recognition)\s*$",
        "interests": r"(?im)^(?:interests?|hobbies|activities|personal\s*interests?)\s*$",
        "languages": r"(?im)^(?:languages?|language\s*proficiency)\s*$",
    }
    
    def __init__(self):
        """Initialize the resume parser."""
        self._check_dependencies()
    
    def _check_dependencies(self):
        """Check and log available parsing libraries."""
        self.available_formats = [".txt"]
        
        if PDF_AVAILABLE or PYPDF2_AVAILABLE:
            self.available_formats.append(".pdf")
        
        if DOCX_AVAILABLE:
            self.available_formats.append(".docx")
    
    def parse(self, file_path: str = None, file_content: bytes = None, 
              file_name: str = None) -> ParsedResume:
        """
        Parse a resume file and extract structured information.
        
        Args:
            file_path: Path to the resume file
            file_content: Raw file content (bytes)
            file_name: Original filename (used for format detection)
            
        Returns:
            ParsedResume object with extracted information
        """
        # Determine file extension
        if file_path:
            ext = Path(file_path).suffix.lower()
            file_name = Path(file_path).name
        elif file_name:
            ext = Path(file_name).suffix.lower()
        else:
            raise ValueError("Either file_path or file_name must be provided")
        
        # Extract raw text based on file format
        if ext == ".pdf":
            raw_text = self._parse_pdf(file_path, file_content)
        elif ext == ".docx":
            raw_text = self._parse_docx(file_path, file_content)
        elif ext == ".txt":
            raw_text = self._parse_txt(file_path, file_content)
        else:
            raise ValueError(f"Unsupported file format: {ext}")
        
        # Create parsed resume object
        resume = ParsedResume(raw_text=raw_text, file_name=file_name or "unknown")
        
        # Extract structured information
        resume.email = self._extract_email(raw_text)
        resume.phone = self._extract_phone(raw_text)
        resume.name = self._extract_name(raw_text)
        resume.location = self._extract_location(raw_text)
        resume.linkedin = self._extract_linkedin(raw_text)
        resume.github = self._extract_github(raw_text)
        resume.skills = self._extract_skills(raw_text)
        resume.education = self._extract_education(raw_text)
        resume.experience = self._extract_experience(raw_text)
        resume.certifications = self._extract_certifications(raw_text)
        resume.projects = self._extract_projects(raw_text)
        resume.publications = self._extract_publications(raw_text)
        resume.awards = self._extract_awards(raw_text)
        resume.interests = self._extract_interests(raw_text)
        resume.languages = self._extract_languages(raw_text)
        resume.summary = self._extract_summary(raw_text)
        resume.total_years_experience = self._calculate_total_experience(resume.experience)
        
        return resume
    
    def _parse_pdf(self, file_path: str = None, file_content: bytes = None) -> str:
        """Extract text from PDF file."""
        text_parts = []
        
        if PDF_AVAILABLE:
            try:
                if file_path:
                    with pdfplumber.open(file_path) as pdf:
                        for page in pdf.pages:
                            page_text = page.extract_text()
                            if page_text:
                                text_parts.append(page_text)
                elif file_content:
                    import io
                    with pdfplumber.open(io.BytesIO(file_content)) as pdf:
                        for page in pdf.pages:
                            page_text = page.extract_text()
                            if page_text:
                                text_parts.append(page_text)
                
                if text_parts:
                    return "\n".join(text_parts)
            except Exception as e:
                print(f"pdfplumber failed: {e}, trying PyPDF2...")
        
        if PYPDF2_AVAILABLE:
            try:
                if file_path:
                    reader = PdfReader(file_path)
                elif file_content:
                    import io
                    reader = PdfReader(io.BytesIO(file_content))
                else:
                    raise ValueError("No file provided")
                
                for page in reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
                
                return "\n".join(text_parts)
            except Exception as e:
                raise RuntimeError(f"Failed to parse PDF: {e}")
        
        raise RuntimeError("No PDF parsing library available. Install pdfplumber or PyPDF2.")
    
    def _parse_docx(self, file_path: str = None, file_content: bytes = None) -> str:
        """Extract text from DOCX file."""
        if not DOCX_AVAILABLE:
            raise RuntimeError("python-docx is not installed. Run: pip install python-docx")
        
        try:
            if file_path:
                doc = Document(file_path)
            elif file_content:
                import io
                doc = Document(io.BytesIO(file_content))
            else:
                raise ValueError("No file provided")
            
            paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
            return "\n".join(paragraphs)
        except Exception as e:
            raise RuntimeError(f"Failed to parse DOCX: {e}")
    
    def _parse_txt(self, file_path: str = None, file_content: bytes = None) -> str:
        """Extract text from TXT file."""
        try:
            if file_path:
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    return f.read()
            elif file_content:
                return file_content.decode('utf-8', errors='ignore')
            else:
                raise ValueError("No file provided")
        except Exception as e:
            raise RuntimeError(f"Failed to parse TXT: {e}")
    
    def _extract_email(self, text: str) -> str:
        """Extract email address from text."""
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        match = re.search(email_pattern, text)
        return match.group(0) if match else ""
    
    def _extract_phone(self, text: str) -> str:
        """Extract phone number from text."""
        phone_patterns = [
            r'\+?1?[-.\s]?\(?[0-9]{3}\)?[-.\s]?[0-9]{3}[-.\s]?[0-9]{4}',
            r'\+?[0-9]{1,3}[-.\s]?[0-9]{3,4}[-.\s]?[0-9]{3,4}[-.\s]?[0-9]{3,4}',
        ]
        for pattern in phone_patterns:
            match = re.search(pattern, text)
            if match:
                return match.group(0)
        return ""
    
    def _extract_name(self, text: str) -> str:
        """Extract candidate name (usually first line or before email)."""
        lines = text.strip().split('\n')
        
        # Name is typically in the first few lines
        for line in lines[:5]:
            line = line.strip()
            # Skip lines that look like headers or contain special characters
            if line and len(line) < 50:
                # Skip if it contains email or phone
                if '@' in line or re.search(r'\d{3}[-.\s]?\d{3}', line):
                    continue
                # Skip common headers
                if any(header in line.lower() for header in ['resume', 'cv', 'curriculum']):
                    continue
                # Return first valid-looking name line
                if re.match(r'^[A-Za-z\s\-\.]+$', line) and len(line) > 2:
                    return line
        
        return ""
    
    def _extract_skills(self, text: str) -> List[str]:
        """Extract skills from resume text."""
        text_lower = text.lower()
        found_skills = []
        
        # Find skills from predefined list
        for skill in self.COMMON_SKILLS:
            if skill.lower() in text_lower:
                # Proper case formatting
                found_skills.append(skill.title() if len(skill) > 3 else skill.upper())
        
        # Also look for skills section and extract additional items
        skills_section = self._extract_section(text, "skills")
        if skills_section:
            # Split by common delimiters
            additional = re.split(r'[,;•|\n]', skills_section)
            for item in additional:
                item = item.strip()
                if item and 2 < len(item) < 30:
                    if item.lower() not in [s.lower() for s in found_skills]:
                        found_skills.append(item)
        
        return list(set(found_skills))[:30]  # Limit to 30 skills
    
    def _extract_section(self, text: str, section_name: str) -> str:
        """Extract content of a specific section."""
        pattern = self.SECTION_PATTERNS.get(section_name)
        if not pattern:
            return ""
        
        # Find section header
        match = re.search(pattern, text)
        if not match:
            return ""
        
        start = match.end()
        
        # Find next section header or end of text
        remaining_text = text[start:]
        next_section = None
        min_pos = len(remaining_text)
        
        for other_pattern in self.SECTION_PATTERNS.values():
            if other_pattern != pattern:
                other_match = re.search(other_pattern, remaining_text)
                if other_match and other_match.start() < min_pos:
                    min_pos = other_match.start()
        
        return remaining_text[:min_pos].strip()
    
    def _extract_education(self, text: str) -> List[Dict[str, str]]:
        """Extract education information with enhanced details."""
        education = []
        edu_section = self._extract_section(text, "education")
        
        if not edu_section:
            edu_section = text
        
        # Split education section into entries by newlines with degree keywords
        lines = edu_section.strip().split('\n')
        
        # Group lines into education entries
        entries = []
        current_entry_lines = []
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            # Check if this line starts a new degree entry
            is_new_entry = bool(re.search(
                r'^(Master|Bachelor|Doctor|Ph\.?D|MBA|M\.?S\.?|B\.?S\.?|B\.?A\.?|M\.?A\.?|Associate)',
                line, re.IGNORECASE
            ))
            
            if is_new_entry and current_entry_lines:
                entries.append('\n'.join(current_entry_lines))
                current_entry_lines = []
            
            current_entry_lines.append(line)
        
        if current_entry_lines:
            entries.append('\n'.join(current_entry_lines))
        
        # Parse each entry
        for entry_text in entries:
            if not entry_text.strip():
                continue
            
            edu_entry = {
                "degree": "",
                "field": "",
                "institution": "",
                "year": "",
                "gpa": "",
                "honors": "",
                "coursework": []
            }
            
            # Extract degree type
            degree_match = re.search(
                r'(Master\s+of\s+Science|Master\s+of\s+Arts|Bachelor\s+of\s+Science|Bachelor\s+of\s+Arts|'
                r'Doctor\s+of\s+Philosophy|Ph\.?D\.?|MBA|M\.?S\.?|B\.?S\.?|B\.?A\.?|M\.?A\.?|Associate)',
                entry_text, re.IGNORECASE
            )
            if degree_match:
                edu_entry["degree"] = degree_match.group(1).strip()
            
            # Extract field of study - look for "in [Field]" pattern
            # Pattern: "Master of Science in Computer Science" -> "Computer Science"
            field_patterns = [
                r'(?:Master|Bachelor|Ph\.?D|MBA|M\.?S\.?|B\.?S\.?)(?:\s+of\s+\w+)?\s+in\s+([A-Za-z][A-Za-z\s]{2,30}?)(?:\s*[-–]|\s*$|\||,|\n|Track|Stanford|University|College)',
                r'[-–]\s*([A-Za-z][A-Za-z\s]{3,25})\s*Track',  # "- Machine Learning Track"
            ]
            for pattern in field_patterns:
                field_match = re.search(pattern, entry_text, re.IGNORECASE)
                if field_match:
                    field = field_match.group(1).strip()
                    # Remove "Science" or "Arts" if they're standalone at start
                    if field.lower().startswith('science in '):
                        field = field[11:]
                    elif field.lower().startswith('arts in '):
                        field = field[8:]
                    field = field.strip()
                    if field and not re.search(r'University|College|Institute|School', field, re.IGNORECASE):
                        edu_entry["field"] = field
                        break
            
            # Extract institution - look for University/College names on a separate line or after location
            # First try to find it on a line with a city/state
            inst_lines = entry_text.split('\n')
            for line in inst_lines:
                inst_match = re.search(
                    r'([A-Z][A-Za-z\s]+(?:University|College|Institute|School))',
                    line
                )
                if inst_match:
                    inst_name = inst_match.group(1).strip()
                    # Make sure we don't include track/concentration text
                    inst_name = re.sub(r'\s*(Track|Concentration|Focus).*', '', inst_name)
                    edu_entry["institution"] = inst_name
                    break
            
            # If not found, try common university names
            if not edu_entry["institution"]:
                common_unis = re.search(
                    r'(Stanford|MIT|Harvard|Berkeley|UCLA|NYU|Columbia|Princeton|Yale|Cornell|CMU|Georgia\s+Tech|Caltech)',
                    entry_text, re.IGNORECASE
                )
                if common_unis:
                    edu_entry["institution"] = common_unis.group(1).strip()
            
            # Extract year range
            year_match = re.search(r'(20\d{2}|19\d{2})\s*[-–]\s*(20\d{2}|19\d{2}|Present|Current)', entry_text, re.IGNORECASE)
            if year_match:
                edu_entry["year"] = f"{year_match.group(1)}-{year_match.group(2)}"
            else:
                # Single year
                single_year = re.search(r'(20\d{2}|19\d{2})', entry_text)
                if single_year:
                    edu_entry["year"] = single_year.group(1)
            
            # Extract GPA
            gpa_match = re.search(r'GPA[:\s]*([0-4]\.\d+)\s*(?:/\s*4\.0)?', entry_text, re.IGNORECASE)
            if gpa_match:
                edu_entry["gpa"] = gpa_match.group(1)
            
            # Extract honors
            honors_match = re.search(
                r'(Summa\s+Cum\s+Laude|Magna\s+Cum\s+Laude|Cum\s+Laude|with\s+(?:highest\s+)?honors?|Dean\'?s?\s+List)',
                entry_text, re.IGNORECASE
            )
            if honors_match:
                edu_entry["honors"] = honors_match.group(1)
            
            # Only add if we found meaningful data
            if edu_entry["degree"] or edu_entry["institution"]:
                education.append(edu_entry)
        
        return education[:5]  # Limit to 5 entries
    
    def _extract_experience(self, text: str) -> List[Dict[str, str]]:
        """Extract work experience information with enhanced details."""
        experience = []
        exp_section = self._extract_section(text, "experience")
        
        if not exp_section:
            exp_section = text
        
        lines = exp_section.split('\n')
        current_job = None
        
        for i, line in enumerate(lines):
            line = line.strip()
            if not line:
                continue
            
            # Check for pipe-separated format: "Title | Company | Location"
            pipe_match = re.match(r'^([^|]+)\|\s*([^|]+)(?:\|\s*(.+))?$', line)
            
            # Check for date/duration patterns on this or next line
            duration_match = re.search(
                r"((?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\.?\s*\d{4})\s*[-–]\s*((?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\.?\s*\d{4}|Present|Current)",
                line, re.IGNORECASE
            )
            year_range_match = re.search(r"(\d{4})\s*[-–]\s*(\d{4}|Present|Current)", line, re.IGNORECASE)
            
            if pipe_match:
                # Save previous job if exists
                if current_job:
                    experience.append(current_job)
                
                title = pipe_match.group(1).strip()
                company = pipe_match.group(2).strip()
                location = pipe_match.group(3).strip() if pipe_match.group(3) else ""
                
                current_job = {
                    "title": title,
                    "company": company,
                    "location": location,
                    "duration": "",
                    "description": "",
                    "achievements": [],
                }
                
                # Check next line for duration
                if i + 1 < len(lines):
                    next_line = lines[i + 1].strip()
                    dur_match = re.search(
                        r"((?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\.?\s*\d{4})\s*[-–]\s*((?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\.?\s*\d{4}|Present|Current)",
                        next_line, re.IGNORECASE
                    )
                    if dur_match:
                        current_job["duration"] = f"{dur_match.group(1)} - {dur_match.group(2)}"
            
            elif duration_match and current_job and not current_job["duration"]:
                # This is a duration line
                current_job["duration"] = f"{duration_match.group(1)} - {duration_match.group(2)}"
            
            elif current_job and len(line) > 10:
                # Check if it's a bullet point (achievement)
                if line.startswith(('•', '-', '*', '▪', '○')) or re.match(r'^\d+\.', line):
                    achievement = line.lstrip('•-*▪○0123456789. ').strip()
                    if achievement:
                        current_job["achievements"].append(achievement)
                elif not year_range_match and not duration_match:
                    # Add to description
                    if current_job["description"]:
                        current_job["description"] += " " + line
                    else:
                        current_job["description"] = line
        
        if current_job:
            experience.append(current_job)
        
        return experience[:10]  # Limit to 10 entries
    
    def _extract_certifications(self, text: str) -> List[Dict[str, str]]:
        """Extract certifications with structured details."""
        certifications = []
        cert_section = self._extract_section(text, "certifications")
        
        # Use cert section if found, otherwise search whole text
        search_text = cert_section if cert_section else text
        
        # Split by bullet points or newlines
        lines = re.split(r'\n\s*[•\-\*]|\n', search_text)
        
        for line in lines:
            line = line.strip().lstrip('•-* ')
            if not line or len(line) < 10:
                continue
            
            # Skip non-certification lines
            if re.match(r'^(EDUCATION|EXPERIENCE|SKILLS|PROJECTS|AWARDS)', line, re.IGNORECASE):
                continue
            
            # Check if line contains certification keywords
            cert_keywords = [
                r'AWS\s+Certified', r'Google\s+Cloud', r'Azure', r'Certified', r'Certificate',
                r'Certification', r'Professional\s+Data', r'Specialization',
                r'TensorFlow', r'Kubernetes', r'CKA', r'CKAD', r'PMP', r'CISSP',
                r'CompTIA', r'Scrum', r'TOGAF', r'ITIL', r'Deep\s+Learning'
            ]
            
            is_cert = any(re.search(kw, line, re.IGNORECASE) for kw in cert_keywords)
            if not is_cert:
                continue
            
            cert_entry = {
                "name": "",
                "issuer": "",
                "year": "",
                "expiry": "",
                "credential_id": ""
            }
            
            # Extract year first (usually at end in parentheses)
            year_match = re.search(r'\(?\s*(20[0-2]\d)\s*\)?', line)
            if year_match:
                cert_entry["year"] = year_match.group(1)
                # Remove year from line for cleaner name extraction
                line_clean = re.sub(r'\s*\(?\s*20[0-2]\d\s*\)?\s*$', '', line)
            else:
                line_clean = line
            
            # Extract the full certification name
            cert_entry["name"] = line_clean.strip()
            
            # Determine issuer based on certification name
            name_lower = cert_entry["name"].lower()
            if 'aws' in name_lower or 'amazon' in name_lower:
                cert_entry["issuer"] = "Amazon Web Services"
            elif 'google' in name_lower or 'gcp' in name_lower:
                cert_entry["issuer"] = "Google Cloud"
            elif 'azure' in name_lower or 'microsoft' in name_lower:
                cert_entry["issuer"] = "Microsoft"
            elif 'tensorflow' in name_lower:
                cert_entry["issuer"] = "Google"
            elif 'coursera' in name_lower:
                cert_entry["issuer"] = "Coursera"
            elif 'kubernetes' in name_lower or 'cka' in name_lower or 'ckad' in name_lower:
                cert_entry["issuer"] = "CNCF"
            elif 'pmp' in name_lower or 'pmi' in name_lower:
                cert_entry["issuer"] = "PMI"
            elif 'scrum' in name_lower:
                cert_entry["issuer"] = "Scrum Alliance"
            elif 'comptia' in name_lower:
                cert_entry["issuer"] = "CompTIA"
            elif 'cisco' in name_lower or 'ccna' in name_lower or 'ccnp' in name_lower:
                cert_entry["issuer"] = "Cisco"
            
            # Check for duplicates
            if cert_entry["name"] and not any(
                c.get('name', '').lower() == cert_entry["name"].lower() 
                for c in certifications
            ):
                certifications.append(cert_entry)
        
        return certifications[:15]
    
    def _extract_summary(self, text: str) -> str:
        """Extract professional summary or objective."""
        summary_section = self._extract_section(text, "summary")
        
        if summary_section:
            # Take first 500 characters
            return summary_section[:500].strip()
        
        # If no explicit summary, use first paragraph
        paragraphs = text.split('\n\n')
        for para in paragraphs:
            para = para.strip()
            if len(para) > 50 and not any(header in para.lower() for header in ['email', 'phone', 'address']):
                return para[:500]
        
        return ""

    def _extract_location(self, text: str) -> str:
        """Extract candidate location."""
        patterns = [
            r"(?i)(?:location|address|based in)[:\s]+([A-Za-z\s,]+(?:,\s*[A-Z]{2})?)",
            r"([A-Z][a-z]+(?:\s+[A-Z][a-z]+)?,\s*[A-Z]{2}(?:\s+\d{5})?)",  # City, STATE ZIP
        ]
        
        for pattern in patterns:
            match = re.search(pattern, text[:1000])
            if match:
                location = match.group(1).strip()
                if 3 < len(location) < 100:
                    return location
        return ""
    
    def _extract_linkedin(self, text: str) -> str:
        """Extract LinkedIn profile URL."""
        pattern = r"(?:linkedin\.com/in/|linkedin:\s*)([A-Za-z0-9\-]+)"
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            return f"linkedin.com/in/{match.group(1)}"
        return ""
    
    def _extract_github(self, text: str) -> str:
        """Extract GitHub profile URL."""
        pattern = r"(?:github\.com/|github:\s*)([A-Za-z0-9\-]+)"
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            return f"github.com/{match.group(1)}"
        return ""
    
    def _extract_projects(self, text: str) -> List[Dict[str, str]]:
        """Extract projects information."""
        projects = []
        proj_section = self._extract_section(text, "projects")
        
        if not proj_section:
            proj_section = text
        
        # Technology keywords to look for
        tech_keywords = [
            'Python', 'Java', 'JavaScript', 'TypeScript', 'React', 'Node', 'Angular', 'Vue',
            'TensorFlow', 'PyTorch', 'Keras', 'Scikit-learn', 'Pandas', 'NumPy',
            'AWS', 'GCP', 'Azure', 'Docker', 'Kubernetes', 'SQL', 'MongoDB', 'Redis',
            'PostgreSQL', 'MySQL', 'Flask', 'Django', 'FastAPI', 'Spring',
            'LSTM', 'CNN', 'NLP', 'BERT', 'GPT', 'Transformer', 'LangChain',
            'Spark', 'Hadoop', 'Kafka', 'Elasticsearch', 'Git', 'MLflow'
        ]
        
        # Split by bullet points
        entries = re.split(r'\n\s*[•\-\*]', proj_section)
        
        for entry in entries:
            entry = entry.strip()
            if not entry or len(entry) < 15:
                continue
            
            project = {
                "name": "",
                "description": "",
                "technologies": [],
                "url": "",
                "highlights": []
            }
            
            # Extract project name - text before colon, dash, or parenthesis
            # Handle formats like:
            # "Project Name - description"
            # "Project Name: description"
            # '"Paper Title" - Venue'
            
            # Check for quoted title (publication style) - handles "curly" and "straight" quotes
            quoted_match = re.search(r'[""\u201c\u201d]([^"""\u201c\u201d]+)[""\u201c\u201d]\s*[-–]\s*(.*)', entry, re.DOTALL)
            if quoted_match:
                project["name"] = quoted_match.group(1).strip()
                project["description"] = quoted_match.group(2).strip()
            else:
                # Try name - description or name: description pattern
                name_match = re.match(r'^([^:\-–\n]{5,60})\s*[:\-–]\s*(.*)', entry, re.DOTALL)
                if name_match:
                    project["name"] = name_match.group(1).strip()
                    project["description"] = name_match.group(2).strip()
                else:
                    # Use first sentence/phrase as name
                    first_part = entry.split('.')[0].split('(')[0].strip()
                    if len(first_part) > 5:
                        project["name"] = first_part[:60]
                        project["description"] = entry
                    else:
                        continue
            
            # Clean up name - remove leading special chars
            project["name"] = re.sub(r'^[•\-\*\s]+', '', project["name"]).strip()
            
            # Extract technologies from the full entry
            for tech in tech_keywords:
                if re.search(rf'\b{re.escape(tech)}\b', entry, re.IGNORECASE):
                    project["technologies"].append(tech)
            
            # Extract URLs
            url_match = re.search(r'(https?://[^\s]+|github\.com/[^\s\)]+)', entry, re.IGNORECASE)
            if url_match:
                project["url"] = url_match.group(1)
            
            # Extract highlights (metrics like stars, users, etc.)
            highlights = re.findall(
                r'(\d+\+?\s*(?:stars|users|downloads|contributors|views|forks))',
                entry, re.IGNORECASE
            )
            project["highlights"] = highlights
            
            # Only add if we have a valid name
            if project["name"] and len(project["name"]) > 3:
                projects.append(project)
        
        return projects[:10]
    
    def _extract_publications(self, text: str) -> List[Dict[str, str]]:
        """Extract publications information."""
        publications = []
        pub_section = self._extract_section(text, "publications")
        
        if not pub_section:
            pub_section = text
        
        # Publication patterns
        pub_patterns = [
            r'"([^"]{20,200})"\s*[-–]\s*([A-Za-z\s]+(?:Conference|Journal|Workshop|Symposium|arXiv|IEEE|ACM)[A-Za-z\s0-9]*)',
            r"([A-Za-z][A-Za-z\s:]+(?:for|using|with|in|of)[A-Za-z\s]+)\s*[-–]\s*((?:KDD|NeurIPS|ICML|ACL|CVPR|EMNLP|ICLR|AAAI|IJCAI)[A-Za-z\s0-9]*)",
            r"(?:published|presented)\s+(?:in|at)\s+([A-Za-z\s]+(?:Conference|Journal|Workshop))",
        ]
        
        for pattern in pub_patterns:
            matches = re.finditer(pattern, pub_section, re.IGNORECASE)
            for match in matches:
                groups = match.groups()
                pub_entry = {
                    "title": groups[0].strip() if groups[0] else "",
                    "venue": groups[1].strip() if len(groups) > 1 and groups[1] else "",
                    "year": "",
                    "authors": "",
                    "url": ""
                }
                
                # Try to extract year
                year_match = re.search(r"20[0-2]\d", match.group(0))
                if year_match:
                    pub_entry["year"] = year_match.group(0)
                
                if pub_entry["title"] and pub_entry not in publications:
                    publications.append(pub_entry)
        
        return publications[:10]
    
    def _extract_awards(self, text: str) -> List[str]:
        """Extract awards and honors."""
        awards = []
        award_section = self._extract_section(text, "awards")
        
        if not award_section:
            award_section = text
        
        # Award patterns
        award_patterns = [
            r"(?i)((?:Best|Outstanding|Excellence|Distinguished|Top)[A-Za-z\s]+(?:Award|Prize|Honor|Recognition))",
            r"(?i)(?:received|awarded|won)\s+(?:the\s+)?([A-Za-z\s]+(?:Award|Prize|Scholarship|Fellowship|Grant))",
            r"(?i)([A-Za-z\s]+(?:Award|Prize|Honor))\s*[-–]\s*(\d{4})",
            r"(?i)(Dean'?s?\s+List|Summa\s+Cum\s+Laude|Magna\s+Cum\s+Laude|Cum\s+Laude|Valedictorian|Salutatorian)",
        ]
        
        for pattern in award_patterns:
            matches = re.finditer(pattern, award_section)
            for match in matches:
                award = match.group(1).strip()
                if award and len(award) > 5 and award not in awards:
                    awards.append(award)
        
        return awards[:10]
    
    def _extract_interests(self, text: str) -> List[str]:
        """Extract personal interests and hobbies."""
        interests = []
        interest_section = self._extract_section(text, "interests")
        
        if not interest_section:
            return interests
        
        # Split by common delimiters
        items = re.split(r'[,;•|\n]', interest_section)
        
        for item in items:
            item = item.strip()
            if item and 2 < len(item) < 50:
                # Filter out common non-interest words
                if not any(word in item.lower() for word in ['experience', 'skill', 'work', 'education']):
                    interests.append(item)
        
        return interests[:15]
    
    def _extract_languages(self, text: str) -> List[str]:
        """Extract language proficiencies."""
        languages = []
        lang_section = self._extract_section(text, "languages")
        
        # Common languages
        common_languages = [
            "english", "spanish", "french", "german", "italian", "portuguese",
            "chinese", "mandarin", "cantonese", "japanese", "korean", "arabic",
            "hindi", "russian", "dutch", "swedish", "norwegian", "danish",
            "polish", "turkish", "greek", "hebrew", "vietnamese", "thai"
        ]
        
        search_text = lang_section if lang_section else text
        text_lower = search_text.lower()
        
        for lang in common_languages:
            if lang in text_lower:
                # Check for proficiency level
                lang_pattern = rf"(?i){lang}\s*[-:–]?\s*(native|fluent|proficient|intermediate|basic|conversational)?"
                match = re.search(lang_pattern, search_text)
                if match:
                    level = match.group(1) if match.group(1) else ""
                    lang_str = lang.title() + (f" ({level.title()})" if level else "")
                    if lang_str not in languages:
                        languages.append(lang_str)
        
        return languages[:10]
    
    def _calculate_total_experience(self, experience: List[Dict[str, str]]) -> float:
        """Calculate total years of experience from work history."""
        total_years = 0.0
        
        for exp in experience:
            duration = exp.get('duration', '')
            
            # Try to extract years from duration
            year_patterns = [
                r"(\d+)\+?\s*years?",
                r"(\d{4})\s*[-–]\s*(\d{4}|present|current)",
                r"(\d+)\s*months?",
            ]
            
            for pattern in year_patterns:
                match = re.search(pattern, duration, re.IGNORECASE)
                if match:
                    groups = match.groups()
                    if len(groups) == 2 and groups[1]:
                        # Year range
                        try:
                            start = int(groups[0])
                            end = 2026 if groups[1].lower() in ['present', 'current'] else int(groups[1])
                            total_years += (end - start)
                        except ValueError:
                            pass
                    elif len(groups) == 1:
                        try:
                            if 'month' in pattern:
                                total_years += int(groups[0]) / 12
                            else:
                                total_years += int(groups[0])
                        except ValueError:
                            pass
                    break
        
        return round(total_years, 1)
